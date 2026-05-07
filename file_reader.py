import re
import PyPDF2
import docx
import os
from docx import Document


def clean_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n+", "\n", text)
    text = re.sub(r"\.\s+", ". ", text)
    return text.strip()


def extract_text(uploaded_file) -> str:
    file_name = uploaded_file.name.lower()

    # PDF
    if file_name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(uploaded_file)
        text_chunks = []

        for page in reader.pages:
            try:
                extracted = page.extract_text()
                if extracted:
                    text_chunks.append(extracted)
            except:
                continue

        return clean_text("\n".join(text_chunks))

    # DOCX
    elif file_name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        text_chunks = []

        for para in doc.paragraphs:
            if para.text:
                text_chunks.append(para.text)

        return clean_text("\n".join(text_chunks))

    return None


def load_resume(path="my_resume/Chaitanya_Kumar_Reddy.docx") -> str:
    if not os.path.exists(path):
        print(f"Resume not found at {path}")
        return None

    doc = Document(path)
    content = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    content.append(cell.text.strip())

    return clean_text("\n".join(content))
