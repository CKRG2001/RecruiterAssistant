from docx import Document
import os


def load_resume(path="my_resume/chaitanya_resume.docx"):
    if not os.path.exists(path):
        print(f"Resume not found at {path}")
        return None

    doc = Document(path)

    content = []

    for para in doc.paragraphs:
        content.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                content.append(cell.text)

    return "\n".join(content)
